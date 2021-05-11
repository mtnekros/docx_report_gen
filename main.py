import re
from io import BytesIO
from datetime import date
from PIL import Image, ImageOps
from docx import Document
from docx.shared import Inches, Twips
from docx.enum.text import WD_BREAK
# from save_charts import generate_all_charts
from report_generation.table import DTable
# generate_all_charts()

# friends_data = [
#     {"id": 1, "name": "Sacheen", "address": "Bhaktapur", "hobbies": "Playing football" },
#     {"id": 2, "name": "Pushpa", "address": "Bhaktapur", "hobbies": "Playing TableTennis" },
#     {"id": 3, "name": "Prism", "address": "Gwarko", "hobbies": "Watching movies" },
#     {"id": 4, "name": "Ashin", "address": "Lalitpur", "hobbies": "Going on adventures" },
# ]

# def add_header_row(table, cols):
#     cells = table.rows[0].cells
#     for index,column in enumerate(cols):
#         cells[index].text = column.get("title")
#         cells[index].paragraphs[0].runs[0].font.bold = True
#         if column.get("width"):
#             cells[index].width = column.get("width")
# 
# def move_table_after(para, table, doc, title="", new_para=False):
#     tbl, p = table._tbl, para._p
#     if title:
#         par2 = doc.add_paragraph(title)
#         par2.runs[0].font.bold = True
#         p.addprevious(par2._p)
#     p.addprevious(tbl)
#     if new_para:
#         par1 = doc.add_paragraph("")
#         p.addprevious(par1._p)
# 
# 
# def generate_friends_table(doc: Document, para):
#     columnDefs = [
#         {"property": "id", "title": "#", "width": Twips(2) },
#         {"property": "name", "title": "Name", "width": Twips(2000) },
#         {"property": "address", "title": "Address", "width": Twips(5000) },
#         {"property": "hobbies", "title": "Hobbies", },
#     ]
#     table = doc.add_table(rows=1, cols=len(columnDefs))
#     table.style = 'Table Grid'
#     move_table_after(para, table, doc, title="List of friends", new_para=False)
#     add_header_row(table, columnDefs)
# 
#     for friend in friends_data:
#         cells = table.add_row().cells
#         for index,cell in enumerate(cells):
#             colInfo = columnDefs[index]
#             cell.text = str(friend.get(colInfo.get("property")))
#             if colInfo.get("width"):
#                 cell.width = colInfo.get("width")


# columnDefs = [
#     {"property": "id", "title": "#", "width": Twips(2) },
#     {"property": "name", "title": "Name", "width": Twips(2000) },
#     {"property": "address", "title": "Address", "width": Twips(5000) },
#     {"property": "hobbies", "title": "Hobbies", },
# ]


def generate_report():
    doc = Document("base/test.docx")
    context = [
        ("<<GEN_YEAR>>", str(date.today().year)),
        ("<<GEN_MONTH>>", str(date.today().strftime("%B"))),
        ("<<PLAN_NAME>>", "Pandav Gupha Municipality Plan, 2030-2040"),
        ("<<MUNICIPALITY>>", "Pandav Gupha Municipality"),
        ("<<AUTHOR_NAME>>", "Diwash Tamang"),
        ("<<FRIENDS_TABLE>>", DTable.getFriendsTable()),
        ("<<EXPENSE_TABLE>>", DTable.getExpenseTable()),
        ("<<BUBBLE_CHART>>", 'charts/BUBBLE_CHART.jpeg'),
        ("<<EA_CHART>>", 'charts/EA_CHART.jpeg'),
        ("<<PAGE_BREAK>>", 'page-break'),
    ]
    print(", ".join([run.text for p in doc.paragraphs for run in p.runs]))
    for (search, replace) in context:
        for para in doc.paragraphs:
            if search not in para.text:
                continue
            for run in para.runs:
                if run.text!=search:
                    continue
                elif re.search("PAGE_BREAK", search):
                    run.text = run.text.replace(search, "")
                    run.add_break(WD_BREAK.PAGE)
                elif re.search("TABLE", search):
                    run.text = run.text.replace(search, "")
                    dtable = replace
                    dtable.add_to(doc, para)
                elif re.search("CHART", search) or re.search("PICTURE", search):
                    #img = Image.open(replace)
                    #img_with_border = ImageOps.expand(img, border=1, fill='black')
                    #imgByteArr = BytesIO()
                    #img_with_border.save(imgByteArr, format='PNG')
                    #replace = imgByteArr
                    run.text = run.text.replace(search, "")
                    run.add_picture(replace, width=Inches(5))
                else:
                    run.text = run.text.replace(search, replace)
                print("Found: \"", search, "\"", replace)

    doc.save('generated_report.docx')

if __name__=="__main__":
    # generate_all_charts()
    generate_report()
