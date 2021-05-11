import warnings
from docx import Document
from docx.shared import Inches, Length
from docx.text.paragraph import Paragraph
from docx.table import Table


def is_list_or_tuple(value):
    return isinstance(value, list) or isinstance(value, tuple)


class DTable:
    def __init__(self, column_defs, row_data, title=None, style=None, new_para=False):
        DTable.validate_data(column_defs, row_data)
        self.title = title
        self.style = style or 'Table Grid'
        self.column_defs = column_defs
        # set allow autofit to true if none of the cols have defined width
        self.allow_autofit = all(col.get('width') is None for col in column_defs)
        self.row_data = row_data
        self.new_para = new_para

    @property
    def n_columns(self):
        return len(self.column_defs)

    @staticmethod
    def validate_data(column_defs, row_data):
        assert is_list_or_tuple(column_defs), "column_defs must be list or tuple"
        assert is_list_or_tuple(row_data), "row_data must be list or tuple"
        assert len(column_defs) > 0, "no of cols must be > 0"
        assert len(row_data) > 0, "no of row data must be > 0"
        assert all(key in col for col in column_defs for key in ("property", "title",)), \
            "all column def must have property, title & width"
        assert all(col.get("property") in row for row in row_data for col in column_defs), \
            "all rows must have property for each of the column"
        # warnings
        table_width = Length(sum(filter(lambda width: width is not None, map(lambda col: col.get('width'), column_defs))))
        if table_width > Inches(6):
            warnings.warn(f'Table width > 6 at {table_width.inches} Inches. (Usually 6 is max width)')

    def add_to(self, doc: Document, para: Paragraph):
        table = doc.add_table(rows=1, cols=self.n_columns)
        table.style = self.style
        self.move_table_before(doc, para, table)
        if not self.allow_autofit:
            self.force_col_width_for_libre(table)
        self.add_header_row(table)

        for row in self.row_data:
            cells = table.add_row().cells
            for index,cell in enumerate(cells):
                column = self.column_defs[index]
                if column.get('width'):
                    cell.width = column.get("width")
                formatter = column.get("formatter", lambda x: x)
                cell.text = str(formatter(row.get(column.get("property"))))
    
    # this is only to needed to make table have fixed width for libre office
    def force_col_width_for_libre(self, table):
        # table.autofit = False
        table.allow_autofit = False
        for col_def, col in zip(self.column_defs, table.columns):
            if col_def.get('width'):
                col.width = col_def.get('width')

    def add_header_row(self, table):
        cells = table.rows[0].cells
        for index,column in enumerate(self.column_defs):
            cells[index].text = column.get("title")
            cells[index].paragraphs[0].runs[0].font.bold = True
            if column.get('width'):
                cells[index].width = column.get("width")

    def move_table_before(self, doc: Document, para: Paragraph, table: Table):
        tbl, p = table._tbl, para._p
        if self.title:
            title_para = doc.add_paragraph(self.title)
            title_para.runs[0].font.bold = True
            p.addprevious(title_para._p)
        p.addprevious(tbl)
        if self.new_para:
            par1 = doc.add_paragraph("")
            p.addprevious(par1._p)

    @staticmethod
    def getFriendsTable():
        row_data = [
            {"id": 1, "name": "Raman", "address": "Bhaktapur", "hobbies": "Playing football", },
            {"id": 2, "name": "Pwoncho", "address": "Bhaktapur", "hobbies": "Playing TableTennis", },
            {"id": 3, "name": "Ainda", "address": "Gwarko", "hobbies": "Watching movies", },
            {"id": 4, "name": "Ratirar", "address": "Lalitpur", "hobbies": "Going on adventures", },
        ]
        column_defs = [
            {"property": "id", "title": "#", "width": Inches(0.5) },
            {"property": "name", "title": "Name", "width": Inches(1) },
            {"property": "address", "title": "Address", "width": Inches(2) },
            {"property": "hobbies", "title": "Hobbies", "width": Inches(2.5) },
        ]
        return DTable(
            column_defs,
            row_data,
            title='Table 1.1.1: Friends Table',
            style='Table Grid'
        )

    @staticmethod
    def getExpenseTable():
        row_data = [
            {"id": 1, "particular": "Food expenses",   "amt":  4000 },
            {"id": 2, "particular": "Travel expenses", "amt":  2000 },
            {"id": 3, "particular": "Wushu fees",      "amt":  1000 },
            {"id": 4, "particular": "Guitar lessons",  "amt":  3000 },
            {"id": 5, "particular": "Social events",   "amt":  2000 },
            {"id": 6, "particular": "Bank payments",   "amt": 10000 },
        ]
        column_defs = [
            {"property": "id",         "title": "#",            "width": Inches(0.5) },
            {"property": "particular", "title": "Particular",   "width": Inches(3) },
            {"property": "amt",        "title": "Amount (Rs.)", "width": Inches(2.5), "formatter": lambda x: f"Rs. {x:,.2f}" },
        ]
        return DTable(
            column_defs,
            row_data,
            title='Table 1.2.1: A summary of my monthly expenses',
            style='Table Grid'
        )
